"""DOCX parser using python-docx."""

import time
import logging
from ..parsers import BaseParser, DocumentParseResult, register_parser

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    extension = ".docx"
    display_name = "DOCX Parser"

    def _import_library(self):
        from docx import Document  # noqa: F401

    def extract_metadata(self, filepath: str) -> dict:
        from docx import Document
        doc = Document(filepath)
        core_props = doc.core_properties
        return {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "author": str(core_props.author) if core_props and core_props.author else None,
            "title": str(core_props.title) if core_props and core_props.title else None,
        }

    def parse(self, filepath: str) -> DocumentParseResult:
        start = time.time()
        try:
            from docx import Document
            doc = Document(filepath)

            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    text_parts.append(" | ".join(cells))

            full_text = "\n\n".join(text_parts)
            word_count = len(full_text.split())
            elapsed = (time.time() - start) * 1000

            logger.debug(
                "Parsed DOCX: %s (%d paragraphs, %d words, %.0fms)",
                filepath, len(doc.paragraphs), word_count, elapsed,
            )
            return DocumentParseResult(
                text=full_text,
                metadata=self.extract_metadata(filepath),
                page_count=0,
                word_count=word_count,
                warnings=[],
                parser_used="docx_parser",
                processing_time_ms=round(elapsed, 2),
            )

        except Exception as e:
            elapsed = (time.time() - start) * 1000
            logger.error("DOCX parsing failed for %s: %s", filepath, e)
            return DocumentParseResult(
                text="",
                metadata={},
                page_count=0,
                word_count=0,
                warnings=[f"DOCX parsing failed: {e}"],
                parser_used="docx_parser",
                processing_time_ms=round(elapsed, 2),
            )


register_parser(DOCXParser())